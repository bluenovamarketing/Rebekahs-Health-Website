from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\todda\Blue Nova Projects\Rebekahs Health Website")
ASSETS = ROOT / "output" / "client-guide" / "assets"
OUT = ROOT / "output" / "client-guide"
DOCX_PATH = OUT / "Rebekahs-Website-Content-Guide.docx"

# compact_reference_guide preset, with a restrained health-and-wellness color override.
GREEN = "285943"
GOLD = "B58A45"
INK = "24312B"
MUTED = "66736D"
PALE = "EEF5F1"
LIGHT_GOLD = "FBF5E9"
WHITE = "FFFFFF"
BODY_FONT = "Calibri"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run(run, size=11, bold=False, italic=False, color=INK, font=BODY_FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)
    return run


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_left_border(paragraph, color=GOLD, size="18", space="8"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)
    pbdr.append(left)


def set_cellish_padding(paragraph):
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.12)


def add_callout(doc, label, text, fill=PALE, accent=GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.15
    set_cellish_padding(p)
    shade_paragraph(p, fill)
    paragraph_left_border(p, accent)
    set_run(p.add_run(label + "  "), size=10.5, bold=True, color=accent)
    set_run(p.add_run(text), size=10.5, color=INK)
    return p


def add_body(doc, text, after=6, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True, color=GREEN)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_run(p.add_run(text))
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    set_run(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], bold=True,
            color=GREEN if level < 3 else INK)
    return p


def set_image_alt(inline_shape, description):
    docpr = inline_shape._inline.docPr
    docpr.set("descr", description)
    docpr.set("title", description)


def add_screenshot(doc, filename, alt, width=6.25):
    path = ASSETS / filename
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_image_alt(shape, alt)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(8)
    set_run(cap.add_run(alt), size=9, italic=True, color=MUTED)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, txt, fld_end])
    set_run(run, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, 18, 10, GREEN),
        "Heading 2": (13, 14, 7, GREEN),
        "Heading 3": (12, 10, 5, INK),
    }
    for name, (size, before, after, color) in heading_specs.items():
        st = styles[name]
        st.font.name = BODY_FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = BODY_FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        st.font.size = Pt(11)
        st.font.color.rgb = rgb(INK)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def add_header_footer(section):
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_run(hp.add_run("REBEKAH'S HEALTH & NUTRITION  |  WEBSITE CONTENT GUIDE"),
            size=8.5, bold=True, color=MUTED)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    set_run(fp.add_run("Client reference  |  Page "), size=9, color=MUTED)
    add_page_number(fp)


def page_break(doc):
    doc.add_page_break()


def build():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    configure_styles(doc)
    add_header_footer(sec)

    # Page 1 - customer-pack style opening.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run("CLIENT WEBSITE GUIDE"), size=10, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    set_run(p.add_run("Create Events & Blog Posts"), size=27, bold=True, color=GREEN)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_run(p.add_run("Rebekah's Health & Nutrition"), size=15, color=MUTED)

    add_callout(doc, "PURPOSE", "Use this guide for everyday content updates. It covers Events and Posts only. Social-feed instructions are maintained separately.", fill=PALE, accent=GREEN)
    add_heading(doc, "The safe publishing routine", 1)
    for item in [
        "Start with a draft. Do not publish until the title, content, date, location, category and featured image are complete.",
        "Use Preview and check the page on both a phone and a computer.",
        "If something looks wrong after publishing, switch it back to Draft and contact Blue Nova Marketing.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Featured-image standard", 2)
    add_body(doc, "Use a clear landscape image in a 16:9 shape. A practical size is 1600 x 900 pixels. JPG or WebP is preferred; keep the file below about 500 KB when practical. Add useful alt text that briefly describes the image.")
    add_callout(doc, "IMPORTANT", "Do not change themes, plugins, caching, WordPress settings, DNS, WooCommerce settings or SEO sitewide settings while creating content.", fill=LIGHT_GOLD, accent=GOLD)

    page_break(doc)
    add_heading(doc, "Create an Event", 1)
    add_body(doc, "Events appear on the public Events calendar and may also appear in event cards elsewhere on the website.")
    add_numbered(doc, "From the WordPress dashboard, choose Events in the left menu.")
    add_numbered(doc, "Select Add New Event at the top of the Events screen.")
    add_numbered(doc, "Work as a draft until every required field is complete.")
    add_screenshot(doc, "events-list-header.png", "Events screen: select Add New Event.", width=6.25)
    add_callout(doc, "DO NOT USE", "The site is not configured for ticket sales or RSVPs. Do not install Event Tickets when WordPress promotes it.", fill=LIGHT_GOLD, accent=GOLD)

    page_break(doc)
    add_heading(doc, "Add the event title and description", 1)
    add_numbered(doc, "Enter a short, specific title, such as 'Nutrition Q&A - Lake Orion'.")
    add_numbered(doc, "In the main editor, explain what the event is, who it is for and what visitors should expect.")
    add_numbered(doc, "Include any preparation instructions, registration method or contact information inside the description.")
    add_numbered(doc, "Choose the matching Event Category: Clarkston, Grand Blanc, Lake Orion or Lapeer.")
    add_screenshot(doc, "event-editor-top.png", "New Event screen: title, description, Publish panel and location categories.", width=6.25)
    add_callout(doc, "WRITING TIP", "Lead with the most useful information. Avoid putting the entire description in bold, all caps or one long paragraph.", fill=PALE, accent=GREEN)

    page_break(doc)
    add_heading(doc, "Set the date, location and image", 1)
    add_numbered(doc, "Under The Events Calendar, set the start date/time and end date/time. Use All Day Event only when no specific time applies.")
    add_numbered(doc, "Leave the site time zone as America/Detroit.")
    add_numbered(doc, "Choose an existing Venue and Organizer when possible. Create a new one only when it is genuinely new.")
    add_numbered(doc, "For free events, leave Cost blank or enter 0. Add an Event Website URL only when visitors need an outside registration or information page.")
    add_numbered(doc, "Select Set featured image and use a 16:9 landscape image.")
    add_screenshot(doc, "event-details.png", "Event details: date and time, venue, organizer, status and featured image.", width=6.25)
    add_callout(doc, "RECURRING EVENTS", "The current calendar does not include the Pro recurrence feature. Create each date separately or duplicate an existing event and carefully update every date-specific field.", fill=LIGHT_GOLD, accent=GOLD)

    page_break(doc)
    add_heading(doc, "Preview, publish and maintain Events", 1)
    add_heading(doc, "Before publishing", 2)
    for item in [
        "Title is clear and contains no spelling errors.",
        "Start and end dates and times are correct.",
        "Correct location category, venue and organizer are selected.",
        "Featured image is landscape and not blurry.",
        "Description explains what, where, when and how to participate.",
        "Preview works on desktop and mobile.",
    ]:
        add_bullet(doc, item)
    add_body(doc, "When everything is correct, select Publish. Open the public Events page afterward and confirm the event appears as expected.")
    add_heading(doc, "Update or cancel an Event", 2)
    add_bullet(doc, "To correct information, open Events, select the event title, make the change and select Update.")
    add_bullet(doc, "For a canceled event, set Events Status to Canceled and update the description with the cancellation message. Do not immediately delete it.")
    add_bullet(doc, "If an event was created by mistake and never announced, move it to Trash only after confirming it is not linked from another page.")
    add_callout(doc, "AFTER ANY CHANGE", "Open the public event in a private/incognito browser window. Check the date, location, image and links.", fill=PALE, accent=GREEN)

    page_break(doc)
    add_heading(doc, "Create a Blog Post", 1)
    add_body(doc, "Blog posts are articles, educational updates and Q&A content. Do not create routine articles as Pages.")
    add_numbered(doc, "From the WordPress dashboard, choose Posts, then Add Post.")
    add_numbered(doc, "Enter a clear headline. Use normal capitalization rather than all caps.")
    add_numbered(doc, "Use the + button to add Paragraph, Heading, Image and List blocks. Use Heading 2 for main sections and Heading 3 only inside those sections.")
    add_numbered(doc, "Use Save draft frequently. Do not select Edit with Elementor for a normal blog article.")
    add_screenshot(doc, "posts-list-header.png", "Posts screen: select Add Post.", width=6.25)
    add_callout(doc, "GOOD ARTICLE STRUCTURE", "Short introduction, helpful sections with headings, concise paragraphs, one clear next step and links to relevant website pages.", fill=PALE, accent=GREEN)

    page_break(doc)
    add_heading(doc, "Complete the Post settings", 1)
    add_numbered(doc, "In the Post sidebar, choose one best-fit Category. Add a new category only when the topic cannot fit an existing one.")
    add_numbered(doc, "Set a 16:9 featured image and add useful alt text.")
    add_numbered(doc, "Write a one- or two-sentence Excerpt that summarizes the article without repeating the full opening paragraph.")
    add_numbered(doc, "To schedule, change Publish from Immediately to the intended date and time. The website uses the America/Detroit time zone.")
    add_screenshot(doc, "blog-editor.png", "Post editor: add blocks in the main area and complete settings in the Post sidebar.", width=6.25)
    add_heading(doc, "Basic SEOPress check", 2)
    add_bullet(doc, "Save the draft first, then open the SEO/No score control in the top toolbar.")
    add_bullet(doc, "Confirm the search title clearly describes the article and includes the main topic naturally.")
    add_bullet(doc, "Write a natural meta description of about one or two sentences. Do not stuff it with repeated keywords.")
    add_bullet(doc, "Leave Advanced, Redirection, Schema and sitewide SEO settings unchanged unless Blue Nova Marketing provides instructions.")

    page_break(doc)
    add_heading(doc, "Final publishing checklist", 1)
    add_heading(doc, "For every Event", 2)
    for item in [
        "Correct title, date, time and America/Detroit time zone",
        "Correct location category, venue and organizer",
        "Clear description and working registration/contact link",
        "16:9 featured image with alt text",
        "Previewed and checked on the public Events page",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "For every Blog Post", 2)
    for item in [
        "Headline, introduction and useful section headings",
        "One best-fit category and a short excerpt",
        "16:9 featured image with alt text",
        "Relevant internal links and basic SEOPress title/description",
        "Previewed on mobile and desktop before publishing",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "When to ask for help", 2)
    add_body(doc, "Contact Blue Nova Marketing before changing themes, plugins, menus, page templates, WooCommerce, site settings, caching, DNS or sitewide SEO. Also ask for help if the editor displays an error, a published item does not appear, or a page looks broken.")
    add_callout(doc, "BEST RECOVERY STEP", "If you are unsure, save the item as a Draft. A draft can be reviewed safely without changing the public website.", fill=LIGHT_GOLD, accent=GOLD)

    props = doc.core_properties
    props.title = "Rebekah's Health & Nutrition - Website Content Guide"
    props.subject = "Client instructions for creating Events and Blog Posts"
    props.author = "Blue Nova Marketing"
    props.keywords = "WordPress, Events, Blog Posts, Rebekah's Health & Nutrition"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
